import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# Page Configuration to Wide Layout to mimic your Canvas design
st.set_page_config(page_title="MANUGTURO AI", page_icon="📝", layout="wide")

# Custom CSS injected to replicate the chalkboard & paper layout perfectly with high contrast
st.markdown("""
    <style>
    .stApp { background-color: #1a261f; }
    [data-testid="stColumn"]:nth-of-type(1) {
        background-color: #e8e4db;
        padding: 25px;
        border-radius: 12px;
        box-shadow: 3px 3px 10px rgba(0,0,0,0.3);
    }
    [data-testid="stColumn"]:nth-of-type(2) {
        background-color: #fdfbf7;
        padding: 30px;
        border-radius: 15px;
        box-shadow: 5px 5px 15px rgba(0,0,0,0.4);
        min-height: 500px;
        border: 1px solid #e0dad0;
    }
    
    /* Targets input labels in the Left Column to force them to pure black text */
    [data-testid="stColumn"]:nth-of-type(1) label p {
        color: #000000 !important;
        font-weight: 600 !important;
    }
    
    /* Fixed text contrast rules for pure black high-visibility readability */
    [data-testid="stColumn"]:nth-of-type(2) div,
    [data-testid="stColumn"]:nth-of-type(2) p,
    [data-testid="stColumn"]:nth-of-type(2) li,
    [data-testid="stColumn"]:nth-of-type(2) span,
    [data-testid="stColumn"]:nth-of-type(2) td,
    [data-testid="stColumn"]:nth-of-type(2) th,
    [data-testid="stColumn"]:nth-of-type(2) ol,
    [data-testid="stColumn"]:nth-of-type(2) ul {
        color: #111111 !important;
        font-weight: 500 !important;
    }
    h1, h5 { color: #ffffff !important; font-family: 'Helvetica Neue', sans-serif; }
    h3 { color: #1a261f !important; font-family: 'Helvetica Neue', sans-serif; }
    
    /* Elegant styling rules for generated markdown tables in the preview */
    table { width: 100%; border-collapse: collapse; margin-top: 15px; margin-bottom: 15px; background-color: #ffffff; }
    th, td { border: 1px solid #1a261f !important; padding: 10px; text-align: left; color: #111111 !important; }
    th { background-color: #e8e4db; font-weight: bold !important; }
    </style>
""", unsafe_allow_html=True)

# App Headers
st.title("MANUGTURO Lesson Planner v.1.0")
st.markdown("##### *Click. Generate. Teach. Custom lessons made easy for FREE...! Developed for DepEd Teachers*")
st.markdown("---")

# EXACT ORIGINAL RETAINED PROMPTS
TEMPLATE_1_ILAW = """
Your task is to generate a highly professional and effective Lesson Plan using the ILAW Framework, following the DepEd Order 9, series of 2026 trimester guidelines and adhere to the Intentions, Learning Experience, Assessment and Ways Forward (ILAW) Framework and its Learning Design Principles (Clear Goals and Teaching, Active Retrieval and Spacing, Self-awareness and Metacognition, Scaffolding, Purpose & Values Integration, Inclusion, Checks for Understanding, and Social Learning).

CRITICAL FORMATTING INSTRUCTION: You MUST use markdown tables (e.g., | Header | Header |) to present structured components, specifically for the Lesson Objectives (split by Cognitive, Psychomotor, Affective domains), the Flow milestones, and the Assessment Accommodations matrix.

Structure the output according to these four pillars:
### 1. INTENTIONS
       * Learning Competency and Curriculum Standards: Write the competency/ies from the curriculum guide that we are targeting, and the content or performance standards applicable to the sessions.            
       * Learning Objectives: Provide specific objectives for Cognitive, Psychomotor, and Affective domains.
       * Learner Context: Describe a realistic classroom context including common learner strengths, interests, and potential barriers to learning (e.g., visual impairments, fine motor difficulties, or varying reading levels).

### 2. LEARNING EXPERIENCE
       * Pre-Lesson: Describe a (get ready) activity that checks well-being and hooks interest.
       * Flow: Design a (journey) of activities that purposefully apply the Learning Design Principles:
               - Make objectives clear to learners.
               - Provide scaffolding (guide them before letting them try alone).
               - Check mastery, understanding, and well-being during the session.
               - Connect new concepts to past competencies.
               - Encourage collaboration and reflection on why the lesson matters.

    * Learning Resources: List inclusive resources and provide (Emergency Alternatives) (e.g., If no internet, use printed diagrams).
    * Opportunities for Integration: Suggest a meaningful connection to another subject or technology.

### 3. ASSESSMENT
       * Formative Assessment: Design 2-3 tasks or questions to gather feedback during the lesson.                  
       * Accommodations: Provide specific ways (visual, auditory, or group-based) for learners with different abilities to demonstrate their understanding.

### 4. WAYS FORWARD
       * Extended Learning: Suggest an activity learners can do at home to spark curiosity or reinforce mastery.
       * Reflections: Provide 2 reflective questions for the teacher to ponder after the session.

### 5. DECLARATION OF AI USE</h3>
       * Include a statement citing how AI was used in the formulation of this plan.
"""

TEMPLATE_4A = """

Create a detailed, learner-centered, and classroom-ready 4A's lesson plan.

CRITICAL FORMATTING INSTRUCTION: You MUST use markdown tables (e.g., | Header | Header |) to present structured components, specifically for the Lesson Objectives (split by Cognitive, Psychomotor, Affective domains), the Flow milestones, and the Assessment Accommodations matrix.

The lesson plan should follow the 4A's instructional model:
    * Activity
    * Analysis
    * Abstraction
    * Application

Include the following components:
    1. Objectives
       Provide:
            * Cognitive objectives
            * Affective objectives
            * Psychomotor objectives

    2. Subject Matter
       Include:
            * Topic
            * References
            * Materials/Visual Aids
            * Subject Integrations

    3. Procedure Proper (4A's)
       A. Activity
          * Motivational activity or review
          * Interactive and collaborative student task
          * Icebreaker or energizer (if appropriate)
          * Teacher instructions and possible student responses

       B. Analysis
          * Processing questions
          * Guided discussion
          * Critical-thinking questions
          * Teacher script and expected learner responses

       C. Abstraction
          * Generalization of the lesson
          * Key concepts and simplified explanations
          * Summary of important ideas

       D. Application
          * Real-life application activity
          * Group work, role play, worksheet, performance task, or situational activity
          * Practical integration of the lesson in daily life

    4. Differentiated Instructions and Activities
       Provide differentiated instruction and activities for:
       Advanced/Fast Learners
          * Enrichment tasks
          * Leadership opportunities
          * Higher-order thinking activities

       Average Learners
          * Guided practice
          * Cooperative learning activities
          * Standard classroom tasks

       Struggling Learners
          * Simplified instructions
          * Scaffolded and step-by-step activities
          * Teacher-guided support

       Learners with Special Educational Needs (if applicable)
          * Modified activities
          * Visual supports/manipulatives
          * Flexible assessment/output options
          * Peer buddy support or accommodations

       Differentiate based on:
          * content
          * process
          * product/output
          * pacing
          * grouping strategies

    5. Assessment
       Include:
          * Formative assessment
          * Performance-based assessment
          * Exit ticket, short quiz, or reflection
          * Rubric or scoring guide

    6. Assignment/Homework
          * Provide a meaningful and relevant assignment connected to the lesson.
                        
    Additional Requirements:
       * Align with DepEd standards and MELCs where applicable
       * Use inclusive, engaging, and age-appropriate strategies
       * Integrate 21st-century skills and values formation
       * Include classroom management strategies
       * Make the lesson practical, interactive, and easy to implement
       * Use clear formatting with headings and bullet points
       * Include suggested teacher scripts and possible student responses throughout the lesson plan.
"""

TEMPLATE_5E = """
Create a detailed, learner-centered, and classroom-ready 5E's lesson plan following the 5E Instructional Model.

CRITICAL FORMATTING INSTRUCTION: Present the core instructional procedures using a structured Markdown Table with columns for: [5E Phase, Time Allotment, Teacher's Script/Activities, Learners' Expected Responses].

The lesson plan should follow the **5E Instructional Model**:
   1. Engage
   2. Explore
   3. Explain
   4. Elaborate
   5. Evaluate

 Include the following parts:
   * Learning Competencies/Objectives
   * Content Standards
   * Performance Standards
   * Most Essential Learning Competencies (MELCs), if applicable
   * Subject Integration (if possible)
   * Values Integration
   * Materials/Resources Needed
   * ICT Integration
   * Classroom Management Strategies

For each phase of the 5E model, provide:
   * Teacher's Activities
   * Learners' Activities
   * Guide Questions
   * Expected Responses
   * Time Allotment
                
   Add **differentiated instruction** for:
      * Fast learners
      * Average learners
      * Struggling learners
      * Learners with special educational needs (if applicable)

 Include varied activities such as:
      * Collaborative/group tasks
      * Hands-on activities
      * Games or interactive tasks
      * Real-life application
      * Critical thinking questions

 In the **Evaluate** section, include:
      * Formative assessment
      * Rubrics or scoring guide
      * Performance task
      * Reflection activity

 The lesson plan should:
      * Be detailed but practical
      * Use learner-centered and inquiry-based strategies
      * Encourage higher-order thinking skills (HOTS)
      * Promote participation and inclusivity
      * Be aligned with DepEd standards and 21st-century skills
"""

# Dictionary Mapping
TEMPLATES_MAP = {
    "ILAW Framework (D.O. 9 s. 2026)": TEMPLATE_1_ILAW,
    "4A's Format": TEMPLATE_4A,
    "5E's Format": TEMPLATE_5E
}

# Helper function to convert text and table blocks into native document tables
def convert_to_docx(text_content):
    doc = Document()
    doc.add_heading('MANUGTURO Generated Lesson Plan', level=1)
    
    lines = text_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('|'):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_line = lines[i].strip()
                if '---' not in table_line:
                    row_cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    table_data.append(row_cells)
                i += 1
            
            if table_data:
                num_rows = len(table_data)
                num_cols = max(len(row) for row in table_data)
                word_table = doc.add_table(rows=num_rows, cols=num_cols)
                word_table.style = 'Table Grid'
                
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_value in enumerate(row):
                        if c_idx < len(word_table.rows[r_idx].cells):
                            word_table.rows[r_idx].cells[c_idx].text = cell_value
            continue
            
        if line.startswith('###'):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        elif line.startswith('*') or line.startswith('-'):
            clean_item = line.lstrip('*').lstrip('-').strip()
            doc.add_paragraph(clean_item, style='List Bullet')
        else:
            if line:
                doc.add_paragraph(line)
        i += 1
                
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Secure API Configuration using Streamlit Secrets Ecosystem
secrets_configured = False
if "GEMINI_KEY" in st.secrets:
    GEMINI_API_KEY = st.secrets["GEMINI_KEY"]
    if GEMINI_API_KEY and not GEMINI_API_KEY.startswith("Please replace"):
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-2.5-flash')
        secrets_configured = True

# SIDE-BY-SIDE LAYOUT
col_input, col_output = st.columns([1, 2])

# LEFT PANEL: INPUT DETAILS
with col_input:
    st.markdown("### **INPUT DETAILS**")
    
    grade = st.selectbox("1. Select Grade Level: *", [f"Grade {i}" for i in range(1, 13)])
    subject = st.text_input("2. Enter Subject Area: *", placeholder="e.g., Mathematics or Science")
    topic = st.text_input("3. Topic: *", placeholder="Specify your Topic here")
    criteria = st.text_area("4. Additional Criteria / Instructions:", placeholder="You can specify additional criteria or instructions here...")
    
    template_choice = st.selectbox("5. Lesson Plan Template: *", list(TEMPLATES_MAP.keys()))
    
    # Conditional logic calculation: checks if required parameters are missing text content
    is_missing_fields = not subject.strip() or not topic.strip()
    
    # Generate button locks/unlocks automatically using the calculated status
    generate_clicked = st.button(
        "Generate Lesson Plan", 
        type="primary", 
        use_container_width=True,
        disabled=is_missing_fields,
        help="Please fill in both the Subject Area and Topic fields to unlock generation." if is_missing_fields else None
    )

# RIGHT PANEL: GENERATOR ENGINE
with col_output:
    if generate_clicked:
        if not secrets_configured:
            st.error("⚠️ Secrets Configuration Error: Could not resolve a valid 'GEMINI_KEY'.")
        else:
            with st.spinner("Manugturo is processing your structured template matrices..."):
                template_rules = TEMPLATES_MAP.get(template_choice, TEMPLATE_1_ILAW)
                
                master_prompt = (
                    f"SYSTEM SETTING: You are an expert Master Teacher and DepEd Curriculum Evaluator. "
                    f"Generate a production-ready, highly detailed lesson plan based on the ruleset parameters. "
                    f"MANDATORY OUTPUT FORMAT RULES:\n"
                    f"1. For the 4A's Format, you MUST convert the Objectives, Procedure Proper (Activity, Analysis, Abstraction, Application), and Differentiated Instructions into structural Markdown Tables.\n"
                    f"2. For the 5E's Format, you MUST convert the core Procedure Proper (Engage, Explore, Explain, Elaborate, Evaluate) and Differentiated Instruction sections into structural Markdown Tables.\n"
                    f"3. All tables must use standard markdown table formatting separating columns with | and headers with |---\n\n"
                    f"METADATA:\nGrade: {grade}\nSubject: {subject}\nTopic: {topic}\nCriteria: {criteria}\n\n"
                    f"RULES:\n{template_rules}"
                )
                
                try:
                    response = model.generate_content(master_prompt)
                    st.success(f"✨ Finalized {template_choice} Draft Generated!")
                    st.markdown("---")
                    
                    st.markdown(response.text, unsafe_allow_html=True)
                    st.markdown("---")
                    
                    docx_data = convert_to_docx(response.text)
                    
                    st.download_button(
                        label="📥 Download Clean Lesson Plan Document (.docx)",
                        data=docx_data,
                        file_name=f"{template_choice.replace(' ', '_')}_{topic.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"AI System Processing Exception Error: {e}")
    else:
        st.info("💡 Complete input setup details inside the left configuration card container and click 'Generate Lesson Plan' to render your master curriculum draft sheet layout here.")
